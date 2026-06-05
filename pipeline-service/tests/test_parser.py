"""
Unit tests for the file parser (Stage 1 of the pipeline).
"""

from __future__ import annotations

import shutil
from pathlib import Path

import pytest

from pipeline.parser import RawText, parse_file


# Skip entire file if optional deps are not installed
DOCX_AVAILABLE = shutil.which("docx2txt") is not None  # not used; python-docx is pure


@pytest.fixture
def tmp_dir(tmp_path: Path) -> Path:
    return tmp_path


class TestTxtParser:
    def test_parses_utf8_txt(self, tmp_dir: Path) -> None:
        path = tmp_dir / "novel.txt"
        path.write_text("第一章\n内容1\n\n第二章\n内容2", encoding="utf-8")

        result = parse_file(path)

        assert isinstance(result, RawText)
        assert "第一章" in result.content
        assert "第二章" in result.content
        assert result.filename == "novel.txt"
        assert result.encoding.lower().startswith("utf")
        assert result.chapter_hints == []

    def test_parses_gbk_txt(self, tmp_dir: Path) -> None:
        path = tmp_dir / "novel_gbk.txt"
        content = "第一章\n内容1\n\n第二章\n内容2"
        path.write_bytes(content.encode("gbk"))

        result = parse_file(path)

        assert "第一章" in result.content
        assert "内容1" in result.content
        assert result.encoding.lower() in ("gbk", "gb2312", "gb18030")

    def test_collapses_excessive_blank_lines(self, tmp_dir: Path) -> None:
        path = tmp_dir / "novel.txt"
        path.write_text("para1\n\n\n\n\npara2", encoding="utf-8")

        result = parse_file(path)

        assert "\n\n\n" not in result.content
        assert "para1" in result.content
        assert "para2" in result.content


class TestMarkdownParser:
    def test_extracts_heading_hints(self, tmp_dir: Path) -> None:
        path = tmp_dir / "novel.md"
        path.write_text(
            "# 第一章 相遇\n\n正文内容\n\n# 第二章 重逢\n\n更多内容",
            encoding="utf-8",
        )

        result = parse_file(path)

        assert result.chapter_hints == ["第一章 相遇", "第二章 重逢"]
        assert "正文内容" in result.content

    def test_supports_markdown_extension(self, tmp_dir: Path) -> None:
        path = tmp_dir / "novel.markdown"
        path.write_text("# 第1章\n\ntext", encoding="utf-8")

        result = parse_file(path)

        assert "第1章" in result.content
        assert "text" in result.content

    def test_extracts_multiple_heading_levels(self, tmp_dir: Path) -> None:
        path = tmp_dir / "novel.md"
        path.write_text(
            "# Act 1\n\n## Chapter 1\n\ntext\n\n## Chapter 2\n\ntext2",
            encoding="utf-8",
        )

        result = parse_file(path)

        assert "Act 1" in result.chapter_hints
        assert "Chapter 1" in result.chapter_hints
        assert "Chapter 2" in result.chapter_hints


class TestDocxParser:
    def test_parses_docx(self, tmp_dir: Path) -> None:
        from docx import Document

        path = tmp_dir / "novel.docx"
        doc = Document()
        doc.add_heading("第一章", level=1)
        doc.add_paragraph("林晓揉了揉眼睛。")
        doc.add_heading("第二章", level=1)
        doc.add_paragraph("闹钟已经响了三遍。")
        doc.save(str(path))

        result = parse_file(path)

        assert "林晓揉了揉眼睛" in result.content
        assert "闹钟已经响了三遍" in result.content
        assert "第一章" in result.chapter_hints
        assert "第二章" in result.chapter_hints
        assert result.encoding == "utf-8"


class TestErrorHandling:
    def test_raises_for_missing_file(self, tmp_dir: Path) -> None:
        with pytest.raises(FileNotFoundError):
            parse_file(tmp_dir / "missing.txt")

    def test_raises_for_unsupported_extension(self, tmp_dir: Path) -> None:
        path = tmp_dir / "novel.pdf"
        path.write_text("dummy", encoding="utf-8")
        with pytest.raises(ValueError, match="Unsupported file format"):
            parse_file(path)
